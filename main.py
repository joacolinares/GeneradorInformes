import openai
import datetime
import argparse
from docx import Document

#client = openai.OpenAI(api_key="Private key de Open Ai")

#Aca seteamos el prompt, se puede modificar para el objetivo que quiera lograr para nuestra App
def obtener_texto_llm(tema):
    respuesta = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Eres un asistente experto en redacción de informes."},
            {"role": "user", "content": f"Escribe un informe detallado sobre {tema} de al menos 300 palabras."}
        ],
        max_tokens=1500  
    )
    contenido = respuesta.choices[0].message.content

    # Si el contenido es muy corto le solicitamos a Open Ai que expanda mas el texto actual. En este caso le pedimos que sea minimo 300, pero podria ser mas o menos. Para probarlo debemos modificar el prompt origial pidiendole menos palabras
    if len(contenido.split()) < 300:
        print(" El texto generado es corto, solicitando más contenido...")
        extra_respuesta = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Continúa el informe con más detalles."},
                {"role": "user", "content": "Añade más información para completar el informe con al menos 300 palabras."}
            ],
            max_tokens=1000
        )
        contenido += "\n\n" + extra_respuesta.choices[0].message.content

    return contenido

# Esta función es la que crea el Word y lo edita para que contenga la respuesta de Open Ai, con el titulo, fecha y el texto mismo. Al final lo guardamos
def crear_documento(tema, contenido):
    doc = Document()

    doc.add_heading(f"Informe sobre: {tema}", level=1)
    
    fecha = datetime.datetime.now().strftime("%d-%m-%Y")
    doc.add_paragraph(f"Fecha de creación: {fecha}")
   
    for parrafo in contenido.split("\n\n"):
        doc.add_paragraph(parrafo)

    tema_limpio = tema.title().replace(" ", " ") 
    nombre_archivo = f"{tema_limpio} {fecha}.docx" 

   
    doc.save(nombre_archivo)
    print(f"Documento guardado con el siguiente nombre: {nombre_archivo}")

# En esta parte nos aseguramos de que el tema seleccionado sea el que se le pase al prompt de Open AI mediante la linea de comandos
if __name__ == "__main__":
    
    parser = argparse.ArgumentParser(description="Generador de informes automáticos con OpenAI")
    parser.add_argument("--tema", required=True, help="Tema del informe")
    
    args = parser.parse_args()
    tema = args.tema

    print(f"Generando informe sobre: {tema}")
    
    contenido = obtener_texto_llm(tema)
    crear_documento(tema, contenido)
